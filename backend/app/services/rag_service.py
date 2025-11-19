from sqlalchemy.orm import Session
from typing import List, Dict, Any, Optional, Tuple
import os
import json
import asyncio
from datetime import datetime
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LangchainDocument
    LANGCHAIN_AVAILABLE = True
except ImportError:
    LANGCHAIN_AVAILABLE = False
    print("⚠️  langchain未安装，RAG功能将受限。请运行: pip install langchain")
    RecursiveCharacterTextSplitter = None
    LangchainDocument = None
from ..models.models import KnowledgeBase, Document, DocumentChunk
from ..models.schemas import (
    KnowledgeBaseCreate, KnowledgeBaseResponse, DocumentUploadRequest,
    RAGSearchRequest, RAGSearchResult
)
from ..services.vector_service import vector_service
from ..services.llm_service import get_llm_service
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    PyPDF2 = None
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    docx = None
try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False
    BeautifulSoup = None
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """文档处理器"""
    def __init__(self):
        if LANGCHAIN_AVAILABLE and RecursiveCharacterTextSplitter:
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
                separators=["\n\n", "\n", " ", ""]
            )
        else:
            self.text_splitter = None

    def extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        if not PYPDF2_AVAILABLE or not PyPDF2:
            logger.warning("PyPDF2未安装，无法提取PDF文本")
            return ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return ""

    def extract_text_from_docx(self, file_path: str) -> str:
        """从DOCX文件提取文本"""
        if not DOCX_AVAILABLE or not docx:
            logger.warning("python-docx未安装，无法提取DOCX文本")
            return ""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""

    def extract_text_from_markdown(self, file_path: str) -> str:
        """从Markdown文件提取文本"""
        if not MARKDOWN_AVAILABLE or not markdown or not BS4_AVAILABLE or not BeautifulSoup:
            # 简单返回原始内容
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            except:
                return ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_text = file.read()
                # 转换为纯文本
                html = markdown.markdown(md_text)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text()
        except Exception as e:
            logger.error(f"Error extracting text from Markdown: {e}")
            return ""

    def extract_text_from_html(self, file_path: str) -> str:
        """从HTML文件提取文本"""
        if not BS4_AVAILABLE or not BeautifulSoup:
            logger.warning("beautifulsoup4未安装，无法提取HTML文本")
            return ""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file, 'html.parser')
                # 移除脚本和样式
                for script in soup(["script", "style"]):
                    script.decompose()
                return soup.get_text()
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {e}")
            return ""

    def extract_text_from_txt(self, file_path: str) -> str:
        """从TXT文件提取文本"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {e}")
            return ""

    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """根据文件类型提取文本"""
        file_type = file_type.lower()

        if file_type == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type in ['doc', 'docx']:
            return self.extract_text_from_docx(file_path)
        elif file_type in ['md', 'markdown']:
            return self.extract_text_from_markdown(file_path)
        elif file_type in ['html', 'htm']:
            return self.extract_text_from_html(file_path)
        elif file_type == 'txt':
            return self.extract_text_from_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return ""

    def split_text_into_chunks(self, text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
        """将文本分割成块"""
        try:
            if not self.text_splitter or not LANGCHAIN_AVAILABLE:
                # 简单的文本分割（不使用langchain）
                words = text.split()
                chunks = []
                current_chunk = []
                current_size = 0
                
                for word in words:
                    if current_size + len(word) + 1 > chunk_size and current_chunk:
                        chunks.append(' '.join(current_chunk))
                        current_chunk = current_chunk[-chunk_overlap:] if chunk_overlap > 0 else []
                        current_size = sum(len(w) for w in current_chunk)
                    current_chunk.append(word)
                    current_size += len(word) + 1
                
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                return chunks if chunks else [text]
            
            # 更新文本分割器的配置
            self.text_splitter.chunk_size = chunk_size
            self.text_splitter.chunk_overlap = chunk_overlap

            # 创建Langchain文档
            langchain_doc = LangchainDocument(page_content=text)

            # 分割文档
            chunks = self.text_splitter.split_documents([langchain_doc])
            return [chunk.page_content for chunk in chunks]
        except Exception as e:
            logger.error(f"Error splitting text into chunks: {e}")
            return [text]  # 如果分割失败，返回整个文本作为单个块

class RAGService:
    """RAG（Retrieval-Augmented Generation）服务"""
    def __init__(self, db: Session):
        self.db = db
        self.document_processor = DocumentProcessor()

    async def create_knowledge_base(self, kb_data: KnowledgeBaseCreate) -> KnowledgeBaseResponse:
        """创建知识库"""
        try:
            knowledge_base = KnowledgeBase(
                name=kb_data.name,
                description=kb_data.description,
                user_id=kb_data.user_id
            )

            self.db.add(knowledge_base)
            self.db.commit()
            self.db.refresh(knowledge_base)

            return KnowledgeBaseResponse.from_orm(knowledge_base)

        except Exception as e:
            self.db.rollback()
            raise Exception(f"Failed to create knowledge base: {e}")

    async def upload_document(self, upload_request: DocumentUploadRequest, file_path: str,
                            original_name: str, file_type: str) -> Dict[str, Any]:
        """上传并处理文档"""
        try:
            # 提取文本
            text_content = self.document_processor.extract_text_from_file(file_path, file_type)

            if not text_content:
                raise Exception("Failed to extract text from document")

            # 分割文本
            chunks = self.document_processor.split_text_into_chunks(
                text_content,
                upload_request.chunk_size,
                upload_request.chunk_overlap
            )

            # 创建文档记录
            filename = os.path.basename(file_path)
            document = Document(
                knowledge_base_id=upload_request.knowledge_base_id,
                filename=filename,
                original_name=original_name,
                file_path=file_path,
                file_type=file_type,
                content=text_content,
                chunk_count=len(chunks)
            )

            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)

            # 创建文档块并生成embeddings
            for i, chunk_content in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=i,
                    metadata={
                        "document_name": original_name,
                        "chunk_size": len(chunk_content),
                        "position": i
                    }
                )

                self.db.add(chunk)
                self.db.commit()
                self.db.refresh(chunk)

                # 添加到向量索引
                await asyncio.to_thread(
                    vector_service.add_document_chunk_embedding,
                    chunk.id, chunk_content, self.db
                )

            return {
                "document_id": document.id,
                "filename": original_name,
                "chunk_count": len(chunks),
                "processing_time": None  # 可以添加处理时间统计
            }

        except Exception as e:
            self.db.rollback()
            raise Exception(f"Failed to upload document: {e}")

    async def search_knowledge_base(self, search_request: RAGSearchRequest) -> List[RAGSearchResult]:
        """搜索知识库"""
        try:
            # 使用向量搜索
            search_results = vector_service.search_similar_documents(
                query=search_request.query,
                limit=search_request.limit,
                threshold=search_request.threshold,
                knowledge_base_ids=search_request.knowledge_base_ids,
                db=self.db
            )

            # 转换为RAGSearchResult
            results = []
            for result in search_results:
                rag_result = RAGSearchResult(
                    content=result["content"],
                    document_id=result["document_id"],
                    chunk_index=result["chunk_index"],
                    score=result["score"],
                    metadata=result.get("meta_data", {})
                )
                results.append(rag_result)

            return results

        except Exception as e:
            logger.error(f"Error searching knowledge base: {e}")
            return []

    async def generate_rag_response(self, query: str, search_results: List[RAGSearchResult],
                                 context: Optional[Dict[str, Any]] = None,
                                 model: str = "gpt-3.5-turbo") -> Dict[str, Any]:
        """生成RAG增强的响应"""
        try:
            # 构建上下文
            context_text = ""
            if search_results:
                context_text = "Relevant information from knowledge base:\n\n"
                for i, result in enumerate(search_results, 1):
                    context_text += f"[{i}] {result.content}\n\n"

            # 构建RAG提示
            rag_prompt = f"""
            You are a helpful assistant with access to a knowledge base. Use the provided context to answer the user's question.

            Context from knowledge base:
            {context_text}

            User Question: {query}

            Instructions:
            1. Use the provided context to inform your answer
            2. If the context doesn't contain relevant information, say so
            3. Cite the source of information using [1], [2], etc.
            4. Be comprehensive and accurate
            5. If you're unsure, acknowledge the limitations

            Additional Context: {json.dumps(context or {})}
            """

            # 调用LLM生成响应
            llm_service = get_llm_service()
            response = await llm_service.chat_completion([
                {"role": "system", "content": "You are a knowledgeable assistant with access to curated information sources."},
                {"role": "user", "content": rag_prompt}
            ], model=model)

            return {
                "response": response["content"],
                "sources_used": len(search_results),
                "context_used": context_text,
                "model_used": model,
                "tokens_used": response.get("tokens_used")
            }

        except Exception as e:
            logger.error(f"Error generating RAG response: {e}")
            return {
                "response": "I apologize, but I encountered an error while processing your request.",
                "sources_used": 0,
                "error": str(e)
            }

    async def get_knowledge_base_stats(self, knowledge_base_id: int) -> Dict[str, Any]:
        """获取知识库统计信息"""
        try:
            # 获取文档数量
            document_count = self.db.query(Document).filter(
                Document.knowledge_base_id == knowledge_base_id
            ).count()

            # 获取文档块数量
            chunk_count = self.db.query(DocumentChunk).filter(
                Document.document_id.in_(
                    self.db.query(Document.id).filter(
                        Document.knowledge_base_id == knowledge_base_id
                    )
                )
            ).count()

            # 获取各种文档类型统计
            document_types = {}
            documents = self.db.query(Document).filter(
                Document.knowledge_base_id == knowledge_base_id
            ).all()

            for doc in documents:
                doc_type = doc.file_type.lower()
                document_types[doc_type] = document_types.get(doc_type, 0) + 1

            return {
                "knowledge_base_id": knowledge_base_id,
                "document_count": document_count,
                "chunk_count": chunk_count,
                "document_types": document_types,
                "last_updated": datetime.utcnow().isoformat()
            }

        except Exception as e:
            logger.error(f"Error getting knowledge base stats: {e}")
            return {}

    async def update_document_index(self, document_id: int):
        """更新文档的向量索引"""
        try:
            document = self.db.query(Document).filter(Document.id == document_id).first()
            if not document:
                return

            # 删除旧的chunks
            self.db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).delete()

            # 重新分割文本
            chunks = self.document_processor.split_text_into_chunks(document.content)

            # 创建新的chunks
            for i, chunk_content in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=document.id,
                    content=chunk_content,
                    chunk_index=i,
                    metadata={
                        "document_name": document.original_name,
                        "chunk_size": len(chunk_content),
                        "position": i,
                        "reindexed_at": datetime.utcnow().isoformat()
                    }
                )

                self.db.add(chunk)
                self.db.commit()
                self.db.refresh(chunk)

                # 添加到向量索引
                await asyncio.to_thread(
                    vector_service.add_document_chunk_embedding,
                    chunk.id, chunk_content, self.db
                )

            # 更新文档信息
            document.chunk_count = len(chunks)
            document.updated_at = datetime.utcnow()
            self.db.commit()

            logger.info(f"Reindexed document {document_id} with {len(chunks)} chunks")

        except Exception as e:
            logger.error(f"Error updating document index: {e}")
            self.db.rollback()

    async def delete_document(self, document_id: int):
        """删除文档"""
        try:
            # 删除文档块
            self.db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).delete()

            # 删除文档
            document = self.db.query(Document).filter(Document.id == document_id).first()
            if document:
                self.db.delete(document)

            self.db.commit()

            # 注意：这里可能需要从向量索引中移除相应的embeddings
            # 这取决于向量服务的具体实现

            logger.info(f"Deleted document {document_id}")

        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            self.db.rollback()

    def get_user_knowledge_bases(self, user_id: Optional[int] = None) -> List[KnowledgeBaseResponse]:
        """获取用户的知识库列表"""
        try:
            query = self.db.query(KnowledgeBase)
            if user_id:
                query = query.filter(KnowledgeBase.user_id == user_id)

            knowledge_bases = query.filter(KnowledgeBase.is_active == True).all()
            return [KnowledgeBaseResponse.from_orm(kb) for kb in knowledge_bases]

        except Exception as e:
            logger.error(f"Error getting user knowledge bases: {e}")
            return []

# 全局RAG服务实例（需要通过依赖注入使用）
def get_rag_service(db: Session) -> RAGService:
    return RAGService(db)